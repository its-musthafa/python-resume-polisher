import google.generativeai as genai
from docx import Document
from keys import GEMINI_API_KEY

# --- Configure Gemini ---
genai.configure(api_key=GEMINI_API_KEY)

# --- Ask user for basic details ---
def get_user_inputs():
    print("🧠 Let's build your resume!\nPlease answer a few quick questions.\n")
    name = input("Your full name: ")
    email = input("Email address: ")
    phone = input("Phone number: ")
    linkedin = input("LinkedIn profile URL (optional): ")
    github = input("GitHub profile URL (optional): ")
    education = input("Highest degree & institution (e.g., B.Sc in CS, MIT): ")
    skills = input("List your technical skills (comma-separated): ")
    experience = input("Briefly describe your past job titles or experience: ")
    goal = input("What's the purpose of this resume? (e.g., Applying for Backend Developer role): ")

    return {
        "name": name,
        "email": email,
        "phone": phone,
        "linkedin": linkedin,
        "github": github,
        "education": education,
        "skills": skills,
        "experience": experience,
        "goal": goal,
    }

# --- Generate resume content using Gemini ---
def generate_resume_content(user_data, model_name="models/gemini-1.5-pro"):
    model = genai.GenerativeModel(model_name=model_name)
    prompt = (
        "Generate a professional, concise resume based on the following inputs. "
        "Use a formal tone and include a Summary, Skills, Experience, and Education sections. "
        "Use markdown formatting (like **bold** for section titles, '* ' for bullets).\n\n"
        f"Name: {user_data['name']}\n"
        f"Email: {user_data['email']}\n"
        f"Phone: {user_data['phone']}\n"
        f"LinkedIn: {user_data['linkedin']}\n"
        f"GitHub: {user_data['github']}\n"
        f"Education: {user_data['education']}\n"
        f"Skills: {user_data['skills']}\n"
        f"Experience: {user_data['experience']}\n"
        f"Career Goal: {user_data['goal']}\n"
    )
    try:
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        print(f"Error generating resume: {e}")
        return ""

# --- Format markdown-style resume to docx ---
def save_resume_to_docx(markdown_text, output_path="generated_resume.docx"):
    doc = Document()
    for line in markdown_text.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        if line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            line = line[2:]
        else:
            p = doc.add_paragraph()

        bold = False
        temp = ""
        i = 0
        while i < len(line):
            if line[i:i+2] == "**":
                if temp:
                    run = p.add_run(temp)
                    run.bold = bold
                    temp = ""
                bold = not bold
                i += 2
            else:
                temp += line[i]
                i += 1
        if temp:
            run = p.add_run(temp)
            run.bold = bold
    doc.save(output_path)
    print(f"✅ Resume saved as {output_path}")

# --- Main ---
if __name__ == "__main__":
    user_inputs = get_user_inputs()
    print("\n🚀 Generating your resume with AI...")
    resume_markdown = generate_resume_content(user_inputs)

    if resume_markdown:
        save_resume_to_docx(resume_markdown)
    else:
        print("❌ Resume generation failed.")
