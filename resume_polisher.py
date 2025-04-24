import docx
import google.generativeai as genai
from keys import GEMINI_API_KEY

genai.configure(api_key=GEMINI_API_KEY)

# --- Read resume (.docx) ---
def read_resume(file_path):
    try:
        doc = docx.Document(file_path)
        full_text = "\n".join([para.text for para in doc.paragraphs])
        return full_text
    except Exception as e:
        print(f"Failed to read resume: {e}")
        return None

# --- Convert markdown with bullet & bold support to docx ---
def write_polished_resume(output_path, polished_text):
    doc = docx.Document()
    for line in polished_text.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        if line.startswith("* "):
            line = line[2:]  # Remove bullet char
            p = doc.add_paragraph(style="List Bullet")
        else:
            p = doc.add_paragraph()

        # Handle bold formatting
        bold = False
        temp = ""
        i = 0
        while i < len(line):
            if line[i:i+2] == "**":
                if temp:
                    run = p.add_run(temp)
                    run.bold = bold
                    temp = ""
                bold = not bold
                i += 2
            else:
                temp += line[i]
                i += 1
        if temp:
            run = p.add_run(temp)
            run.bold = bold

    try:
        doc.save(output_path)
        print(f"✅ Polished resume saved to: {output_path}")
    except Exception as e:
        print(f"❌ Failed to save polished resume: {e}")

# --- Send to Gemini for polishing ---
def polish_resume_with_gemini(text, company, designation, model_name="models/gemini-1.5-pro"):
    model = genai.GenerativeModel(model_name=model_name)
    prompt = (
        f"You are a professional resume enhancement assistant. Polish the following resume content to optimize it for the job role "
        f"of '{designation}' at '{company}'. Ensure the formatting is clean and ATS-friendly. Enhance the tone and structure while "
        f"keeping the layout intact. Use markdown formatting (like **bold** for headings, '* ' for bullet points).\n\n"
        f"Here is the resume:\n{text}"
    )
    try:
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        print(f"❌ Error while polishing resume: {e}")
        return ""

# --- Score resume ---
def score_resume(text, model_name="models/gemini-1.5-pro"):
    model = genai.GenerativeModel(model_name=model_name)
    prompt = (
        "Evaluate the following resume on a scale from 0 to 100 based on professionalism, clarity, ATS optimization, and formatting. "
        "Just respond with the number only:\n\n" + text
    )
    try:
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        print(f"❌ Error scoring resume: {e}")
        return "Error"

# --- MAIN ---
if __name__ == "__main__":
    input_path = "generated_resume.docx"
    output_path = "polished_resume.docx"
    model_to_use = "models/gemini-1.5-pro"

    print("📄 Reading original resume...")
    original_text = read_resume(input_path)

    if original_text:
        print("\n🔍 Let's tailor your resume for a job!")
        company = input("🏢 Company you're applying to: ")
        designation = input("💼 Job title/designation: ")

        print("\n✨ Polishing resume with Gemini...")
        polished_text = polish_resume_with_gemini(original_text, company, designation, model_to_use)

        if polished_text:
            print("📝 Writing polished resume to file...")
            write_polished_resume(output_path, polished_text)

            print("\n📊 Scoring resumes...")
            original_score = score_resume(original_text, model_to_use)
            polished_score = score_resume(polished_text, model_to_use)

            print(f"\n✅ Resume Score Comparison:\n- Original: {original_score}/100\n- Polished: {polished_score}/100")
        else:
            print("❌ Resume polishing failed.")
    else:
        print("❌ Could not read the original resume.")
